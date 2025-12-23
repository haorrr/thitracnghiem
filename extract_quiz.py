import json
import sys
from docx import Document
import re

def extract_quiz_from_docx(file_path):
    """Extract quiz questions from a Word document where correct answers are underlined."""
    try:
        doc = Document(file_path)
        questions = []
        current_question = None
        current_answers = []
        correct_answer_index = -1

        # Process all paragraphs
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()

            if not text:
                i += 1
                continue

            # Check if this is a question (starts with "Câu" followed by number)
            if text.startswith("Câu") or (re.match(r'^\d+[:.)\s]', text)):
                # Save previous question if exists
                if current_question and current_answers:
                    questions.append({
                        "question": current_question,
                        "answers": current_answers,
                        "correctAnswer": correct_answer_index
                    })

                # Start new question
                current_question = text
                current_answers = []
                correct_answer_index = -1

                # Look ahead for answers in the same or following paragraphs
                j = i
                answer_count = 0

                # Check if answers are in the same paragraph (inline)
                inline_answers = re.findall(r'([A-D])[.:)\s]+([^A-D]+?)(?=[A-D][.:)\s]|$)', text)

                if inline_answers:
                    # Answers are inline in the question
                    for idx, (letter, answer_text) in enumerate(inline_answers):
                        answer_text = answer_text.strip()
                        if answer_text:
                            current_answers.append(answer_text)

                            # Check if this answer is underlined
                            for run in para.runs:
                                if run.underline or (run.font.underline and answer_text in run.text):
                                    correct_answer_index = idx
                                    break
                else:
                    # Look for answers in following paragraphs
                    j = i + 1
                    while j < len(doc.paragraphs):
                        next_para = doc.paragraphs[j]
                        next_text = next_para.text.strip()

                        if not next_text:
                            j += 1
                            continue

                        # Stop if we hit another question
                        if next_text.startswith("Câu") or re.match(r'^\d+[:.)\s]', next_text):
                            break

                        # Check if this is an answer
                        answer_match = re.match(r'^([A-D])[.:)\s]+(.+)', next_text, re.IGNORECASE)
                        if answer_match:
                            answer_text = answer_match.group(2).strip()
                            current_answers.append(answer_text)

                            # Check if this answer is underlined
                            has_underline = False
                            for run in next_para.runs:
                                if run.underline or run.font.underline:
                                    has_underline = True
                                    break

                            if has_underline:
                                correct_answer_index = answer_count

                            answer_count += 1
                            j += 1
                        else:
                            # Not an answer, might be continuation of previous text
                            j += 1
                            break

                    i = j - 1  # Move to the last processed paragraph

            i += 1

        # Don't forget the last question
        if current_question and current_answers:
            questions.append({
                "question": current_question,
                "answers": current_answers,
                "correctAnswer": correct_answer_index
            })

        return questions

    except Exception as e:
        print(f"Error reading document: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return []

if __name__ == "__main__":
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer)

    file_path = r"D:\Downloads\Thitracnghiem\cauhoi.docx"

    questions = extract_quiz_from_docx(file_path)

    if questions:
        # Output as formatted JSON
        print(json.dumps(questions, ensure_ascii=False, indent=2))
    else:
        print("[]")
